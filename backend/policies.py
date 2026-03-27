"""
policies.py — All policy API routes with SSE streaming support.
"""

import json
import html
import logging
import os
import tempfile
from datetime import datetime, timezone, timedelta
from typing import Generator, Literal, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.background import BackgroundTasks
from pydantic import BaseModel

from .auth import verify_token
from .config import POLICY_LIST
from .database import get_db
from .agents import (
    stream_chat_response,
    stream_generate_policy,
    stream_review_policy,
    stream_finalize_policy,
    stream_revise_policy,
    stream_renew_policy,
)
from .audit import log_event, EVENT_POLICY_GENERATED, EVENT_POLICY_REVIEWED, EVENT_EXPORT_DOCX, EVENT_EXPORT_PDF
from .validators import sanitize_text_input

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/admin/policies", tags=["policies"])


# ── Request / Response models ────────────────────────────────────────────────

class ChatMessage(BaseModel):
    role: Literal["user", "assistant"]
    content: str


class ChatRequest(BaseModel):
    question: str
    history: list[ChatMessage] = []


class PolicyRequest(BaseModel):
    request: str


class PolicyDraft(BaseModel):
    policy_text: str


class ReviseRequest(BaseModel):
    policy_text: str
    comments: str


class FinalizeRequest(BaseModel):
    policy_text: str
    doc_id: str = ""
    title: str = ""
    issue: str = "01"
    rev: str = "00"
    date: str = ""
    prepared_by: str = ""
    reviewed_by: str = ""
    approved_by: str = ""
    department: str = "Information Technology"


# ── SSE helper ───────────────────────────────────────────────────────────────

def _sse_stream(generator: Generator) -> Generator[str, None, None]:
    """
    Wraps a text generator into SSE format.
    Each chunk is emitted as:   data: {"chunk": "..."}
    Terminates with:            data: [DONE]
    """
    try:
        for chunk in generator:
            if chunk:
                payload = json.dumps({"chunk": chunk}, ensure_ascii=False)
                yield f"data: {payload}\n\n"
    except Exception as exc:
        logger.error("SSE stream error: %s", exc)
        error_payload = json.dumps({"chunk": "\n\n[An error occurred while processing your request.]"})
        yield f"data: {error_payload}\n\n"
    finally:
        yield "data: [DONE]\n\n"


# ── Routes ───────────────────────────────────────────────────────────────────

@router.get("/list")
def list_policies(current_user: str = Depends(verify_token)):
    """Return the full list of corporate policies, merging the hardcoded library
    with any additional policies uploaded or saved to managed_policies."""
    all_titles = list(POLICY_LIST)
    known = set(t.lower() for t in all_titles)
    try:
        with get_db() as conn:
            rows = conn.execute("SELECT title FROM managed_policies ORDER BY title").fetchall()
            for row in rows:
                title = row["title"]
                if title.lower() not in known:
                    all_titles.append(title)
                    known.add(title.lower())
    except Exception:
        pass
    return {"policies": all_titles, "count": len(all_titles)}


@router.post("/ingest")
def ingest_policies(current_user: str = Depends(verify_token)):
    """Re-ingest all policy PDFs into ChromaDB (admin only)."""
    from .rag import ingest_documents
    count = ingest_documents()
    return {"status": "ok", "chunks_ingested": count}


@router.post("/chat/stream")
def chat_stream(req: ChatRequest, current_user: str = Depends(verify_token)):
    """SSE stream: AEGIS expert chat response."""
    question = sanitize_text_input(req.question, max_length=2000)
    history_dicts = [{"role": m.role, "content": sanitize_text_input(m.content, max_length=3000)} for m in req.history]
    generator = stream_chat_response(question, history_dicts)
    return StreamingResponse(
        _sse_stream(generator),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/generate/stream")
def generate_stream(req: PolicyRequest, current_user: str = Depends(verify_token)):
    """SSE stream: policy writer agent generates a full policy document."""
    sanitized_request = sanitize_text_input(req.request, max_length=5000)
    log_event(EVENT_POLICY_GENERATED, username=current_user, detail=sanitized_request[:200])
    generator = stream_generate_policy(sanitized_request)
    return StreamingResponse(
        _sse_stream(generator),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/review/stream")
def review_stream(req: PolicyDraft, current_user: str = Depends(verify_token)):
    """SSE stream: compliance auditor agent reviews a policy draft."""
    sanitized_text = sanitize_text_input(req.policy_text, max_length=50000)
    log_event(EVENT_POLICY_REVIEWED, username=current_user)
    generator = stream_review_policy(sanitized_text)
    return StreamingResponse(
        _sse_stream(generator),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


class DocxExportRequest(BaseModel):
    policy_text: str
    doc_id: str = "ASH-IT-POL-000"
    title: str = ""
    issue: str = "01"
    rev: str = "00"
    date: str = ""
    prepared_by: str = "Information Technology Department"
    reviewed_by: str = "Shumon A Zaman"
    approved_by: str = "Shamis Al Dhaheri"
    department: str = "Information Technology"


@router.post("/export/docx")
def export_docx(
    req: DocxExportRequest,
    background_tasks: BackgroundTasks,
    current_user: str = Depends(verify_token),
):
    """
    Export policy draft as an ASH-formatted DOCX matching the corporate template:
    Cover metadata table, revision history table, signature block,
    numbered sections with List Paragraph style, and End of Section marker.
    """
    req.policy_text = sanitize_text_input(req.policy_text, max_length=50000)
    try:
        import re as _re
        import docx as docx_module
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = docx_module.Document()

        # ── Page setup ────────────────────────────────────────────
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Derive title from first line if not provided
        lines = req.policy_text.strip().split("\n")
        first_line = lines[0].lstrip("#").strip() if lines else "IT Policy Document"
        doc_title = req.title.strip() or first_line
        doc_date = req.date or datetime.now(timezone.utc).strftime("%d.%m.%Y")
        initials = lambda name: "".join(w[0] for w in name.split() if w).upper()[:4]

        # ── Helper: set cell shading ──────────────────────────────
        def set_cell_shading(cell, color_hex):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), color_hex)
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

        # ── Helper: format a table cell ───────────────────────────
        def fmt_cell(cell, text, bold=False, size=10, color=None, align=None):
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
            if align:
                p.alignment = align

        # ═══ PAGE 1: Title ════════════════════════════════════════
        doc.add_paragraph("")  # spacer

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(doc_title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        # Org subtitle
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run("Ali & Sons Holding LLC")
        sr.font.size = Pt(12)
        sr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        sub2 = doc.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr2 = sub2.add_run("DIH IT Security Division — Abu Dhabi, UAE")
        sr2.font.size = Pt(10)
        sr2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        doc.add_paragraph("")

        # ── Metadata table (ASH format: 6 rows x 4 cols) ─────────
        meta_table = doc.add_table(rows=6, cols=4)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.style = "Table Grid"

        meta_data = [
            ("Prepared by:", req.prepared_by, "Document Status:", "For Implementation"),
            ("Reviewed by:", req.reviewed_by, "Identification Number:", req.doc_id),
            ("Approved by:", req.approved_by, "Issue:", req.issue),
            ("Responsible Department:", req.department, "Revision:", req.rev),
            ("Format:", "A4", "Language:", "English"),
            ("Date:", doc_date, "Classification:", "INTERNAL"),
        ]
        for r_idx, (l1, v1, l2, v2) in enumerate(meta_data):
            row = meta_table.rows[r_idx]
            fmt_cell(row.cells[0], l1, bold=True, size=9, color=(0x47, 0x55, 0x69))
            set_cell_shading(row.cells[0], "F8FAFC")
            fmt_cell(row.cells[1], v1, size=9.5)
            fmt_cell(row.cells[2], l2, bold=True, size=9, color=(0x47, 0x55, 0x69))
            set_cell_shading(row.cells[2], "F8FAFC")
            fmt_cell(row.cells[3], v2, size=9.5, bold=(r_idx == 1))  # bold doc ID

        doc.add_paragraph("")

        # ── Revision History table ────────────────────────────────
        rev_heading = doc.add_paragraph()
        rr = rev_heading.add_run("Revision History")
        rr.font.size = Pt(12)
        rr.font.bold = True

        rev_table = doc.add_table(rows=2, cols=7)
        rev_table.style = "Table Grid"
        rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rev_headers = ["Issue", "Revision", "Date", "Description", "Prepared by", "Reviewed by", "Approved by"]
        for ci, h in enumerate(rev_headers):
            cell = rev_table.rows[0].cells[ci]
            fmt_cell(cell, h, bold=True, size=9, color=(0xFF, 0xFF, 0xFF))
            set_cell_shading(cell, "0F172A")

        # First revision row
        rev_row = rev_table.rows[1]
        rev_vals = [req.issue, req.rev, doc_date, "Initial issue", initials(req.prepared_by), initials(req.reviewed_by), initials(req.approved_by)]
        for ci, v in enumerate(rev_vals):
            fmt_cell(rev_row.cells[ci], v, size=9)

        doc.add_paragraph("")

        # ── Signature block table ─────────────────────────────────
        sig_table = doc.add_table(rows=3, cols=3)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_roles = [("Prepared by", req.prepared_by), ("Reviewed by", req.reviewed_by), ("Approved by", req.approved_by)]
        for ci, (role, name) in enumerate(sig_roles):
            fmt_cell(sig_table.rows[0].cells[ci], role, bold=True, size=9, color=(0x47, 0x55, 0x69), align=WD_ALIGN_PARAGRAPH.CENTER)
            fmt_cell(sig_table.rows[1].cells[ci], "____________________", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            fmt_cell(sig_table.rows[2].cells[ci], name, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # ═══ PAGE BREAK → Policy Content ═════════════════════════
        doc.add_page_break()

        # ── Parse markdown into sections and render ───────────────
        body_lines = lines[1:] if len(lines) > 1 else lines

        for line in body_lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # H1: # Title (skip — already on cover page)
            if stripped.startswith("# "):
                continue

            # H2: ## Section heading → Bold "List Paragraph"-style heading
            if stripped.startswith("## "):
                heading_text = stripped[3:].strip()
                h_para = doc.add_paragraph()
                h_para.space_before = Pt(12)
                h_para.space_after = Pt(4)
                h_run = h_para.add_run(heading_text)
                h_run.font.size = Pt(13)
                h_run.font.bold = True
                h_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                continue

            # H3: ### Sub-section heading
            if stripped.startswith("### "):
                heading_text = stripped[4:].strip()
                h_para = doc.add_paragraph()
                h_para.space_before = Pt(8)
                h_run = h_para.add_run(heading_text)
                h_run.font.size = Pt(11)
                h_run.font.bold = True
                h_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                continue

            # Bullet list items
            if stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:]
                p = doc.add_paragraph(style="List Bullet")
                # Handle bold within bullet text
                parts = _re.split(r"(\*\*[^*]+\*\*)", text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2])
                        r.font.bold = True
                        r.font.size = Pt(10)
                    else:
                        r = p.add_run(part)
                        r.font.size = Pt(10)
                continue

            # Numbered list items (4.1, 4.1.1, etc.)
            num_match = _re.match(r"^(\d+(?:\.\d+)*)\s+(.+)", stripped)
            if num_match:
                num_str, text = num_match.group(1), num_match.group(2)
                p = doc.add_paragraph()
                p.space_before = Pt(4)
                nr = p.add_run(f"{num_str}\t")
                nr.font.size = Pt(10)
                nr.font.bold = True
                nr.font.color.rgb = RGBColor(0x10, 0xD9, 0xA0)
                # Handle bold in text
                parts = _re.split(r"(\*\*[^*]+\*\*)", text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2])
                        r.font.bold = True
                        r.font.size = Pt(10)
                    else:
                        r = p.add_run(part)
                        r.font.size = Pt(10)
                continue

            # Horizontal rule
            if _re.match(r"^---+$", stripped):
                hr = doc.add_paragraph()
                hr.space_before = Pt(6)
                hr.space_after = Pt(6)
                continue

            # Regular paragraph — handle inline bold
            p = doc.add_paragraph()
            parts = _re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                    r.font.size = Pt(10)
                else:
                    r = p.add_run(part)
                    r.font.size = Pt(10)

        # ── End of Section marker ─────────────────────────────────
        doc.add_paragraph("")
        end_para = doc.add_paragraph()
        end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_run = end_para.add_run("End of Section")
        end_run.font.size = Pt(11)
        end_run.font.bold = True
        end_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # ── Footer note ──────────────────────────────────────────
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer_para.add_run("Generated by AEGIS — AI-Powered IT Policy Manager for Ali & Sons Holding")
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        fr.font.italic = True

        # Save to temp file
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(tmp_path)

        background_tasks.add_task(os.unlink, tmp_path)
        log_event(EVENT_EXPORT_DOCX, username=current_user)

        safe_title = _re.sub(r"[^a-zA-Z0-9_\- ]", "", doc_title).replace(" ", "_")[:50]
        return FileResponse(
            path=tmp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{req.doc_id}_{safe_title}.docx",
        )
    except Exception as exc:
        logger.error("DOCX export error: %s", exc)
        raise HTTPException(status_code=500, detail="DOCX export failed. Please try again.")


@router.post("/export/pdf")
def export_pdf(
    req: DocxExportRequest,
    background_tasks: BackgroundTasks,
    current_user: str = Depends(verify_token),
):
    """
    Export policy draft as a styled PDF.
    Uses fpdf2 (pure Python — no external binary needed).
    """
    req.policy_text = sanitize_text_input(req.policy_text, max_length=50000)
    try:
        from fpdf import FPDF

        class PolicyPDF(FPDF):
            def header(self):
                self.set_font("Helvetica", "B", 10)
                self.set_text_color(26, 86, 219)
                self.cell(0, 5, "Ali & Sons Holding", new_x="LMARGIN", new_y="NEXT")
                self.set_font("Helvetica", "", 8)
                self.set_text_color(107, 114, 128)
                self.cell(0, 4, "DIH IT Security Division  |  Abu Dhabi, UAE", new_x="LMARGIN", new_y="NEXT")
                self.cell(0, 4, "Classification: INTERNAL  |  Generated by AEGIS Policy Manager", new_x="LMARGIN", new_y="NEXT")
                self.set_draw_color(26, 86, 219)
                self.set_line_width(0.8)
                self.line(self.l_margin, self.get_y() + 3, self.w - self.r_margin, self.get_y() + 3)
                self.ln(8)

            def footer(self):
                self.set_y(-20)
                self.set_draw_color(229, 231, 235)
                self.set_line_width(0.3)
                self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                self.ln(3)
                self.set_font("Helvetica", "I", 7)
                self.set_text_color(156, 163, 175)
                self.cell(0, 4, "Generated by AEGIS - AI-Powered IT Policy Manager for Ali & Sons Holding.", align="C", new_x="LMARGIN", new_y="NEXT")
                self.cell(0, 4, "Validate against current UAE NESA and ISO 27001:2022 standards before official publication.", align="C")

        pdf = PolicyPDF(orientation="P", unit="mm", format="A4")
        pdf.set_auto_page_break(auto=True, margin=25)
        pdf.set_margins(left=20, top=20, right=20)
        pdf.add_page()

        lines = req.policy_text.strip().split("\n")

        for line in lines:
            stripped = line.strip()

            if not stripped:
                pdf.ln(3)
                continue

            # H1
            if stripped.startswith("# "):
                text = stripped[2:].strip()
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(26, 86, 219)
                pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            # H2
            elif stripped.startswith("## "):
                text = stripped[3:].strip()
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(30, 58, 138)
                pdf.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
                pdf.set_draw_color(229, 231, 235)
                pdf.set_line_width(0.2)
                pdf.line(pdf.l_margin, pdf.get_y() + 1, pdf.w - pdf.r_margin, pdf.get_y() + 1)
                pdf.ln(3)
            # H3
            elif stripped.startswith("### "):
                text = stripped[4:].strip()
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(30, 64, 175)
                pdf.cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
            # Bullet
            elif stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:].strip()
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(26, 32, 44)
                pdf.cell(6, 5, "-", new_x="END")
                pdf.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")
            # Numbered list
            elif len(stripped) > 1 and stripped[0].isdigit() and stripped[1] in (".", ")"):
                num_part = stripped[:2]
                text = stripped[2:].strip()
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(26, 32, 44)
                pdf.cell(8, 5, num_part, new_x="END")
                pdf.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")
            # Regular paragraph
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(26, 32, 44)
                pdf.multi_cell(0, 5, stripped, new_x="LMARGIN", new_y="NEXT")

        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        pdf.output(tmp_path)
        background_tasks.add_task(os.unlink, tmp_path)
        log_event(EVENT_EXPORT_PDF, username=current_user)

        return FileResponse(
            path=tmp_path,
            media_type="application/pdf",
            filename="AliAndSons_IT_Policy.pdf",
        )
    except Exception as exc:
        logger.error("PDF export error: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="PDF export failed. Please try again.",
        )


@router.post("/generate/revise/stream")
def revise_stream(req: ReviseRequest, current_user: str = Depends(verify_token)):
    """SSE stream: revise a policy draft based on user comments/feedback."""
    sanitized_text = sanitize_text_input(req.policy_text, max_length=50000)
    sanitized_comments = sanitize_text_input(req.comments, max_length=5000)
    log_event(EVENT_POLICY_GENERATED, username=current_user, detail=f"Revise: {sanitized_comments[:100]}")
    generator = stream_revise_policy(sanitized_text, sanitized_comments)
    return StreamingResponse(
        _sse_stream(generator),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/generate/finalize/stream")
def finalize_stream(req: FinalizeRequest, current_user: str = Depends(verify_token)):
    """SSE stream: finalize an approved policy draft into publication-ready format."""
    sanitized_text = sanitize_text_input(req.policy_text, max_length=50000)
    metadata = {
        "Document ID": req.doc_id,
        "Title": req.title,
        "Issue": req.issue,
        "Revision": req.rev,
        "Date": req.date,
        "Prepared By": req.prepared_by,
        "Reviewed By": req.reviewed_by,
        "Approved By": req.approved_by,
        "Department": req.department,
    }
    log_event(EVENT_POLICY_GENERATED, username=current_user, detail=f"Finalize: {req.title[:100]}")
    generator = stream_finalize_policy(sanitized_text, metadata)
    return StreamingResponse(
        _sse_stream(generator),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/renew/stream")
def renew_stream(req: PolicyDraft, current_user: str = Depends(verify_token)):
    """SSE stream: renewal agent reviews and updates an existing policy for annual renewal."""
    sanitized_text = sanitize_text_input(req.policy_text, max_length=50000)
    log_event(EVENT_POLICY_REVIEWED, username=current_user, detail="Annual renewal review")
    from .agents import stream_renew_policy
    generator = stream_renew_policy(sanitized_text)
    return StreamingResponse(
        _sse_stream(generator),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/cache-status")
def policy_cache_status(current_user: str = Depends(verify_token)):
    """Return the current state of the policy cache."""
    from .policy_cache import get_cache_status
    return get_cache_status()


# ── Managed policies models ─────────────────────────────────────────────────

class ManagedPolicyCreate(BaseModel):
    title: str
    doc_id: str = ""
    policy_markdown: str = ""
    department: str = "Information Technology"


class ManagedPolicyUpdate(BaseModel):
    title: Optional[str] = None
    doc_id: Optional[str] = None
    policy_markdown: Optional[str] = None
    status: Optional[str] = None
    department: Optional[str] = None


# ── Managed policies endpoints ───────────────────────────────────────────────

@router.get("/managed/renewals")
def list_renewal_policies(current_user: str = Depends(verify_token)):
    """Return managed policies that are due for renewal."""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    try:
        with get_db() as conn:
            rows = conn.execute(
                "SELECT id, title, doc_id, status, version, created_at, updated_at, "
                "next_renewal, renewed_at, department FROM managed_policies "
                "WHERE next_renewal <= ? OR status = 'renewal_due' "
                "ORDER BY next_renewal ASC",
                (today,),
            ).fetchall()
            policies = [dict(row) for row in rows]
    except Exception:
        policies = []
    return {"policies": policies, "count": len(policies)}


# NOTE: Fixed-path routes MUST come before {policy_id} routes to avoid path conflicts.

class SaveDraftRequest(BaseModel):
    title: str
    policy_markdown: str
    doc_id: str = ""
    department: str = "Information Technology"
    audit_result: str = ""


@router.post("/managed/save-draft")
def save_draft_to_library(
    body: SaveDraftRequest,
    current_user: str = Depends(verify_token),
):
    """
    Save a generated policy draft into the managed policy library.
    Creates a new entry with status 'draft' and stores both the policy
    markdown and optional audit result.
    """
    title = sanitize_text_input(body.title, max_length=500).strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")

    policy_md = sanitize_text_input(body.policy_markdown, max_length=50000)
    if not policy_md.strip():
        raise HTTPException(status_code=400, detail="Policy content is required")

    doc_id = sanitize_text_input(body.doc_id, max_length=200).strip()
    department = sanitize_text_input(body.department, max_length=200)
    audit = sanitize_text_input(body.audit_result, max_length=50000)

    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    next_renewal = (datetime.now(timezone.utc) + timedelta(days=365)).strftime("%Y-%m-%d")

    with get_db() as conn:
        if not doc_id:
            max_row = conn.execute(
                "SELECT doc_id FROM managed_policies ORDER BY id DESC LIMIT 1"
            ).fetchone()
            next_num = 1
            if max_row and max_row["doc_id"]:
                import re
                m = re.search(r"(\d+)$", max_row["doc_id"])
                if m:
                    next_num = int(m.group(1)) + 1
            doc_id = f"POL-{next_num:03d}"

        cursor = conn.execute(
            "INSERT INTO managed_policies "
            "(title, doc_id, status, policy_markdown, last_audit, version, created_at, updated_at, "
            "next_renewal, created_by, department) "
            "VALUES (?, ?, 'draft', ?, ?, 1, ?, ?, ?, ?, ?)",
            (title, doc_id, policy_md, audit[:10000], now, now, next_renewal, current_user, department),
        )
        new_id = cursor.lastrowid

    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from .rag import get_vector_store
        from langchain_core.documents import Document
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        doc = Document(page_content=policy_md, metadata={"source": title, "doc_id": doc_id})
        chunks = splitter.split_documents([doc])
        store = get_vector_store()
        store.add_documents(chunks)
    except Exception as exc:
        logger.warning("SaveDraft: ChromaDB ingestion failed (non-fatal): %s", exc)

    log_event(EVENT_POLICY_GENERATED, username=current_user, detail=f"Saved to library: {title}")
    return {"id": new_id, "doc_id": doc_id, "title": title, "status": "draft"}


@router.post("/managed/{policy_id}/revalidate/stream")
def revalidate_managed_policy(
    policy_id: int,
    current_user: str = Depends(verify_token),
):
    """
    SSE stream: Run a fresh compliance review on a managed policy.
    Updates the policy status to 'under_review' while running,
    and stores the audit result in last_audit.
    If no stored markdown, pulls from the policy cache automatically.
    """
    with get_db() as conn:
        row = conn.execute(
            "SELECT id, title, policy_markdown FROM managed_policies WHERE id = ?",
            (policy_id,),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Managed policy not found")

        policy_data = dict(row)
        policy_text = policy_data["policy_markdown"]

        # Backfill from policy cache if empty
        if not policy_text or len(policy_text.strip()) < 50:
            try:
                from .policy_cache import get_policy_summary
                cached = get_policy_summary(policy_data["title"])
                if cached and cached.get("full_text"):
                    policy_text = cached["full_text"]
                    conn.execute(
                        "UPDATE managed_policies SET policy_markdown = ? WHERE id = ?",
                        (policy_text[:50000], policy_id),
                    )
            except Exception:
                pass

        if not policy_text or len(policy_text.strip()) < 50:
            raise HTTPException(status_code=400, detail="Policy has no content to review. Upload the policy document first.")

        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        conn.execute(
            "UPDATE managed_policies SET status = 'under_review', updated_at = ? WHERE id = ?",
            (now, policy_id),
        )

    log_event(EVENT_POLICY_REVIEWED, username=current_user, detail=f"Revalidate: {policy_data['title']}")

    audit_chunks: list[str] = []

    def _revalidate_generator():
        generator = stream_review_policy(policy_text)
        for chunk in generator:
            audit_chunks.append(chunk)
            yield chunk

    def _save_audit_result():
        full_audit = "".join(audit_chunks)
        try:
            with get_db() as conn:
                conn.execute(
                    "UPDATE managed_policies SET last_audit = ?, updated_at = ? WHERE id = ?",
                    (full_audit[:10000], datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"), policy_id),
                )
        except Exception as exc:
            logger.error("Revalidate: failed to save audit result: %s", exc)

    def _sse_with_save():
        try:
            for chunk in _revalidate_generator():
                if chunk:
                    payload = json.dumps({"chunk": chunk}, ensure_ascii=False)
                    yield f"data: {payload}\n\n"
        except Exception as exc:
            logger.error("Revalidate SSE error: %s", exc)
            error_payload = json.dumps({"chunk": "\n\n[Error during revalidation.]"})
            yield f"data: {error_payload}\n\n"
        finally:
            yield "data: [DONE]\n\n"
            _save_audit_result()

    return StreamingResponse(
        _sse_with_save(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.get("/managed/{policy_id}")
def get_managed_policy(policy_id: int, current_user: str = Depends(verify_token)):
    """Return a single managed policy with full markdown content."""
    with get_db() as conn:
        row = conn.execute(
            "SELECT id, title, doc_id, status, policy_markdown, version, "
            "created_at, updated_at, next_renewal, renewed_at, created_by, department "
            "FROM managed_policies WHERE id = ?",
            (policy_id,),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Managed policy not found")
    return {"policy": dict(row)}


@router.get("/managed")
def list_managed_policies(current_user: str = Depends(verify_token)):
    """Return all managed policies (without full markdown body)."""
    try:
        with get_db() as conn:
            rows = conn.execute(
                "SELECT id, title, doc_id, status, version, created_at, updated_at, "
                "next_renewal, renewed_at, department FROM managed_policies "
                "ORDER BY updated_at DESC"
            ).fetchall()
            policies = [dict(row) for row in rows]
    except Exception:
        policies = []
    return {"policies": policies, "count": len(policies)}


@router.post("/managed")
def create_managed_policy(body: ManagedPolicyCreate, current_user: str = Depends(verify_token)):
    """Create a new managed policy entry."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    next_renewal = (datetime.now(timezone.utc) + timedelta(days=365)).strftime("%Y-%m-%d")

    title = sanitize_text_input(body.title, max_length=500)
    doc_id = sanitize_text_input(body.doc_id, max_length=200)
    policy_markdown = sanitize_text_input(body.policy_markdown, max_length=50000)
    department = sanitize_text_input(body.department, max_length=200)

    if not title.strip():
        raise HTTPException(status_code=400, detail="Title is required")

    with get_db() as conn:
        cursor = conn.execute(
            "INSERT INTO managed_policies "
            "(title, doc_id, status, policy_markdown, version, created_at, updated_at, "
            "next_renewal, created_by, department) "
            "VALUES (?, ?, 'approved', ?, 1, ?, ?, ?, ?, ?)",
            (title, doc_id, policy_markdown, now, now, next_renewal, current_user, department),
        )
        new_id = cursor.lastrowid

    return {"id": new_id, "status": "created"}


@router.put("/managed/{policy_id}")
def update_managed_policy(policy_id: int, body: ManagedPolicyUpdate, current_user: str = Depends(verify_token)):
    """Update an existing managed policy."""
    with get_db() as conn:
        row = conn.execute(
            "SELECT id, status, version FROM managed_policies WHERE id = ?",
            (policy_id,),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Managed policy not found")

        existing = dict(row)
        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        updates = ["updated_at = ?"]
        params: list = [now]

        if body.title is not None:
            updates.append("title = ?")
            params.append(sanitize_text_input(body.title, max_length=500))
        if body.doc_id is not None:
            updates.append("doc_id = ?")
            params.append(sanitize_text_input(body.doc_id, max_length=200))
        if body.policy_markdown is not None:
            updates.append("policy_markdown = ?")
            params.append(sanitize_text_input(body.policy_markdown, max_length=50000))
        if body.department is not None:
            updates.append("department = ?")
            params.append(sanitize_text_input(body.department, max_length=200))
        if body.status is not None:
            new_status = sanitize_text_input(body.status, max_length=50)
            valid_statuses = {"draft", "under_review", "approved", "renewal_due", "renewing"}
            if new_status not in valid_statuses:
                raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {', '.join(sorted(valid_statuses))}")
            updates.append("status = ?")
            params.append(new_status)
            # If status changes to approved, set next_renewal to 1 year from now and bump version
            if new_status == "approved" and existing["status"] != "approved":
                next_renewal = (datetime.now(timezone.utc) + timedelta(days=365)).strftime("%Y-%m-%d")
                updates.append("next_renewal = ?")
                params.append(next_renewal)
                updates.append("version = ?")
                params.append(existing["version"] + 1)

        params.append(policy_id)
        conn.execute(
            f"UPDATE managed_policies SET {', '.join(updates)} WHERE id = ?",
            params,
        )

    return {"id": policy_id, "status": "updated"}


@router.post("/managed/{policy_id}/renew")
def renew_managed_policy(policy_id: int, current_user: str = Depends(verify_token)):
    """Trigger renewal for a managed policy."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    with get_db() as conn:
        row = conn.execute(
            "SELECT id FROM managed_policies WHERE id = ?",
            (policy_id,),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Managed policy not found")

        conn.execute(
            "UPDATE managed_policies SET status = 'renewing', renewed_at = ?, updated_at = ? WHERE id = ?",
            (now, now, policy_id),
        )

    return {"id": policy_id, "status": "renewing"}


# ── Upload policy (PDF or DOCX) ──────────────────────────────────────────────

_ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def _extract_text_from_docx(file_path: str) -> tuple[str, int]:
    """Extract all text from a DOCX file. Returns (full_text, page_count)."""
    import docx as docx_module
    doc = docx_module.Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    full_text = "\n\n".join(paragraphs)
    # Approximate page count (DOCX doesn't store page count explicitly)
    page_estimate = max(1, len(full_text) // 3000)
    return full_text, page_estimate


@router.post("/upload")
async def upload_policy(
    file: UploadFile = File(...),
    title: str = Form(""),
    department: str = Form("Information Technology"),
    current_user: str = Depends(verify_token),
):
    """
    Upload a PDF or DOCX policy document.
    - Extracts text from PDF (PyPDFLoader) or DOCX (python-docx)
    - Ingests chunks into ChromaDB for RAG retrieval
    - Creates a managed_policies entry with the extracted text
    Returns the new managed policy ID.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in _ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Accepted formats: PDF, DOCX.",
        )

    if file.size and file.size > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Maximum 20 MB.")

    # Save to a temp file
    fd, tmp_path = tempfile.mkstemp(suffix=ext)
    try:
        file_content = await file.read()
        os.write(fd, file_content)
        os.close(fd)

        # Extract text based on file type
        full_text = ""
        page_count = 0

        if ext == ".pdf":
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(tmp_path)
            pages = loader.load()
            full_text = "\n\n".join(p.page_content for p in pages)
            page_count = len(pages)
        elif ext == ".docx":
            full_text, page_count = _extract_text_from_docx(tmp_path)

        if not full_text or len(full_text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail=f"Could not extract text from {ext.upper()} file. File may be empty or image-only.",
            )

        # Ingest into ChromaDB
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from langchain_core.documents import Document as LCDocument
        from .rag import get_vector_store

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

        if ext == ".pdf":
            # PDF: use the page-level documents for better metadata
            chunks = splitter.split_documents(pages)
        else:
            # DOCX: wrap the full text in a LangChain Document
            lc_doc = LCDocument(
                page_content=full_text,
                metadata={"source": file.filename, "type": "docx"},
            )
            chunks = splitter.split_documents([lc_doc])

        store = get_vector_store()
        store.add_documents(chunks)
        chunks_added = len(chunks)
        logger.info("Upload: ingested %d chunks from '%s' (%s) into ChromaDB", chunks_added, file.filename, ext)

        # Derive title from filename if not provided
        policy_title = sanitize_text_input(title, max_length=500).strip()
        if not policy_title:
            policy_title = os.path.splitext(file.filename)[0].replace("_", " ").replace("-", " ").strip()

        # Create managed policy entry
        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        next_renewal = (datetime.now(timezone.utc) + timedelta(days=365)).strftime("%Y-%m-%d")
        safe_dept = sanitize_text_input(department, max_length=200)

        with get_db() as conn:
            max_row = conn.execute(
                "SELECT doc_id FROM managed_policies ORDER BY id DESC LIMIT 1"
            ).fetchone()
            next_num = 1
            if max_row and max_row["doc_id"]:
                import re
                m = re.search(r"(\d+)$", max_row["doc_id"])
                if m:
                    next_num = int(m.group(1)) + 1
            doc_id = f"POL-{next_num:03d}"

            cursor = conn.execute(
                "INSERT INTO managed_policies "
                "(title, doc_id, status, policy_markdown, version, created_at, updated_at, "
                "next_renewal, created_by, department) "
                "VALUES (?, ?, 'draft', ?, 1, ?, ?, ?, ?, ?)",
                (policy_title, doc_id, full_text[:50000], now, now, next_renewal, current_user, safe_dept),
            )
            new_id = cursor.lastrowid

        log_event(EVENT_POLICY_GENERATED, username=current_user, detail=f"Uploaded: {policy_title} ({ext})")

        return {
            "id": new_id,
            "doc_id": doc_id,
            "title": policy_title,
            "status": "draft",
            "chunks_ingested": chunks_added,
            "pages": page_count,
            "word_count": len(full_text.split()),
            "file_type": ext[1:],  # "pdf" or "docx"
        }

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Upload failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Upload processing failed: {str(exc)}")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ── Batch upload (multiple files) ────────────────────────────────────────────

@router.post("/upload/batch")
async def upload_policies_batch(
    files: list[UploadFile] = File(...),
    department: str = Form("Information Technology"),
    current_user: str = Depends(verify_token),
):
    """
    Upload multiple PDF/DOCX policy documents at once.
    Each file is processed independently — failures on one file don't block others.
    Returns a results array with per-file status.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided.")
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Maximum 20 files per batch.")

    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LCDocument
    from .rag import get_vector_store

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    store = get_vector_store()
    safe_dept = sanitize_text_input(department, max_length=200)
    results = []

    for file in files:
        fname = file.filename or "unknown"
        ext = os.path.splitext(fname)[1].lower()

        if ext not in _ALLOWED_EXTENSIONS:
            results.append({"filename": fname, "status": "error", "detail": f"Unsupported format '{ext}'"})
            continue

        if file.size and file.size > 20 * 1024 * 1024:
            results.append({"filename": fname, "status": "error", "detail": "File too large (>20 MB)"})
            continue

        fd, tmp_path = tempfile.mkstemp(suffix=ext)
        try:
            content = await file.read()
            os.write(fd, content)
            os.close(fd)

            # Extract text
            full_text = ""
            page_count = 0
            pages = None

            if ext == ".pdf":
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(tmp_path)
                pages = loader.load()
                full_text = "\n\n".join(p.page_content for p in pages)
                page_count = len(pages)
            elif ext == ".docx":
                full_text, page_count = _extract_text_from_docx(tmp_path)

            if not full_text or len(full_text.strip()) < 50:
                results.append({"filename": fname, "status": "error", "detail": "Could not extract text"})
                continue

            # Ingest into ChromaDB
            if ext == ".pdf" and pages:
                chunks = splitter.split_documents(pages)
            else:
                lc_doc = LCDocument(page_content=full_text, metadata={"source": fname, "type": ext[1:]})
                chunks = splitter.split_documents([lc_doc])

            store.add_documents(chunks)

            # Create managed policy entry
            policy_title = os.path.splitext(fname)[0].replace("_", " ").replace("-", " ").strip()
            now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
            next_renewal = (datetime.now(timezone.utc) + timedelta(days=365)).strftime("%Y-%m-%d")

            with get_db() as conn:
                max_row = conn.execute("SELECT doc_id FROM managed_policies ORDER BY id DESC LIMIT 1").fetchone()
                next_num = 1
                if max_row and max_row["doc_id"]:
                    import re
                    m = re.search(r"(\d+)$", max_row["doc_id"])
                    if m:
                        next_num = int(m.group(1)) + 1
                doc_id = f"POL-{next_num:03d}"

                cursor = conn.execute(
                    "INSERT INTO managed_policies "
                    "(title, doc_id, status, policy_markdown, version, created_at, updated_at, "
                    "next_renewal, created_by, department) "
                    "VALUES (?, ?, 'draft', ?, 1, ?, ?, ?, ?, ?)",
                    (policy_title, doc_id, full_text[:50000], now, now, next_renewal, current_user, safe_dept),
                )
                new_id = cursor.lastrowid

            results.append({
                "filename": fname,
                "status": "ok",
                "id": new_id,
                "doc_id": doc_id,
                "title": policy_title,
                "chunks_ingested": len(chunks),
                "pages": page_count,
                "word_count": len(full_text.split()),
                "file_type": ext[1:],
            })
            logger.info("Batch upload: processed '%s' → %d chunks", fname, len(chunks))

        except Exception as exc:
            logger.warning("Batch upload: failed on '%s': %s", fname, exc)
            results.append({"filename": fname, "status": "error", "detail": str(exc)[:200]})
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    ok_count = sum(1 for r in results if r["status"] == "ok")
    log_event(EVENT_POLICY_GENERATED, username=current_user, detail=f"Batch upload: {ok_count}/{len(results)} files")
    return {"results": results, "total": len(results), "succeeded": ok_count}


# ── Revalidate and update a managed policy ───────────────────────────────────

@router.post("/managed/{policy_id}/revalidate-update/stream")
def revalidate_and_update_policy(
    policy_id: int,
    current_user: str = Depends(verify_token),
):
    """
    SSE stream: Two-phase revalidation.
    Phase 1 — Compliance audit with latest web research
    Phase 2 — AI rewrites the policy incorporating audit findings + latest standards
    Saves both the audit result and the updated policy to the database.
    If the managed policy has no stored markdown, it pulls content from the
    policy cache (PDF-extracted text) automatically.
    """
    with get_db() as conn:
        row = conn.execute(
            "SELECT id, title, doc_id, policy_markdown, version FROM managed_policies WHERE id = ?",
            (policy_id,),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Managed policy not found")

        policy_data = dict(row)
        policy_text = policy_data["policy_markdown"]

        # If no stored markdown, try to pull from policy cache (PDF text)
        if not policy_text or len(policy_text.strip()) < 50:
            try:
                from .policy_cache import get_policy_summary
                cached = get_policy_summary(policy_data["title"])
                if cached and cached.get("full_text"):
                    policy_text = cached["full_text"]
                    # Also backfill the DB so future calls don't need this lookup
                    conn.execute(
                        "UPDATE managed_policies SET policy_markdown = ? WHERE id = ?",
                        (policy_text[:50000], policy_id),
                    )
                    logger.info("Backfilled policy_markdown for '%s' from cache (%d chars)", policy_data["title"], len(policy_text))
            except Exception as exc:
                logger.warning("Failed to pull from cache: %s", exc)

        if not policy_text or len(policy_text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Policy has no content to review. Upload the policy document first via the Upload tab.",
            )

        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        conn.execute(
            "UPDATE managed_policies SET status = 'under_review', updated_at = ? WHERE id = ?",
            (now, policy_id),
        )

    log_event(EVENT_POLICY_REVIEWED, username=current_user, detail=f"Revalidate-update: {policy_data['title']}")

    from .agents import stream_review_policy, stream_renew_policy

    audit_chunks: list[str] = []
    update_chunks: list[str] = []

    def _stream_both_phases():
        # Phase 1: Compliance audit
        phase_msg = json.dumps({"phase": "audit", "chunk": ""}, ensure_ascii=False)
        yield f"data: {phase_msg}\n\n"

        try:
            for chunk in stream_review_policy(policy_text):
                if chunk:
                    audit_chunks.append(chunk)
                    payload = json.dumps({"phase": "audit", "chunk": chunk}, ensure_ascii=False)
                    yield f"data: {payload}\n\n"
        except Exception as exc:
            logger.error("Revalidate-update audit error: %s", exc)
            err = json.dumps({"phase": "audit", "chunk": f"\n[Audit error: {exc}]"})
            yield f"data: {err}\n\n"

        # Phase 2: Generate updated policy using renewal agent (includes web research)
        phase_msg = json.dumps({"phase": "update", "chunk": ""}, ensure_ascii=False)
        yield f"data: {phase_msg}\n\n"

        try:
            for chunk in stream_renew_policy(policy_text):
                if chunk:
                    update_chunks.append(chunk)
                    payload = json.dumps({"phase": "update", "chunk": chunk}, ensure_ascii=False)
                    yield f"data: {payload}\n\n"
        except Exception as exc:
            logger.error("Revalidate-update renewal error: %s", exc)
            err = json.dumps({"phase": "update", "chunk": f"\n[Update error: {exc}]"})
            yield f"data: {err}\n\n"

        yield "data: [DONE]\n\n"

        # Save results to DB
        full_audit = "".join(audit_chunks)
        full_update = "".join(update_chunks)
        try:
            with get_db() as conn:
                updates = [
                    "last_audit = ?",
                    "updated_at = ?",
                    "status = 'approved'",
                    "version = ?",
                ]
                params = [
                    full_audit[:10000],
                    datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
                    policy_data["version"] + 1,
                ]
                # If renewal agent produced an updated policy, save it
                if full_update and len(full_update.strip()) > 200:
                    updates.append("policy_markdown = ?")
                    params.append(full_update[:50000])
                    # Reset renewal date
                    updates.append("next_renewal = ?")
                    params.append((datetime.now(timezone.utc) + timedelta(days=365)).strftime("%Y-%m-%d"))

                params.append(policy_id)
                conn.execute(
                    f"UPDATE managed_policies SET {', '.join(updates)} WHERE id = ?",
                    params,
                )
                logger.info(
                    "Revalidate-update: saved audit (%d chars) + updated policy (%d chars) for %s",
                    len(full_audit), len(full_update), policy_data["title"],
                )
        except Exception as exc:
            logger.error("Revalidate-update DB save failed: %s", exc)

    return StreamingResponse(
        _stream_both_phases(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
